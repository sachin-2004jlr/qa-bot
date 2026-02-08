import streamlit as st
import os
import shutil
import uuid
import io
from docx import Document
from src.backend import AdvancedRAG

# 1. Page Configuration
st.set_page_config(
    page_title="Multi Model RAG", 
    layout="wide",
    initial_sidebar_state="expanded"
)

# 2. Advanced Professional Dark Theme (CSS)
st.markdown("""
    <style>
    /* Global App Background */
    .stApp {
        background-color: #0e1117;
    }

    /* Main Title Styling - Refined Modern Serif/Sans */
    .main-title {
        text-align: center;
        font-family: 'Inter', -apple-system, sans-serif;
        color: #f3f4f6;
        padding: 40px 0px 20px 0px;
        font-weight: 700;
        letter-spacing: -0.05em;
        text-transform: uppercase;
        border-bottom: 1px solid #1f2937;
        margin-bottom: 30px;
    }
    
    /* Sidebar Professional Styling */
    section[data-testid="stSidebar"] {
        background-color: #111827;
        border-right: 1px solid #1f2937;
    }
    
    section[data-testid="stSidebar"] .stMarkdown h2, 
    section[data-testid="stSidebar"] .stMarkdown h1 {
        color: #9ca3af;
        font-size: 0.85rem;
        text-transform: uppercase;
        letter-spacing: 0.1em;
        margin-top: 20px;
    }

    /* Chat Container General Styling */
    .chat-container {
        padding: 24px;
        border-radius: 8px;
        margin-bottom: 20px;
        border: 1px solid #1f2937;
        font-family: 'Inter', sans-serif;
    }
    
    /* User Message Styling - Shady Black with Slate Blue Accent */
    .user-box {
        background-color: #1f2937;
        border-left: 4px solid #475569; /* Slate Blue-Grey accent */
    }
    
    /* AI Message Styling - Dark Charcoal with Deep Emerald Accent */
    .ai-box {
        background-color: #111827;
        border-left: 4px solid #065f46; /* Deep Emerald accent */
        box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.5);
    }
    
    /* Metadata/Role Header Styling */
    .role-header {
        font-weight: 600;
        color: #6b7280;
        margin-bottom: 12px;
        text-transform: uppercase;
        font-size: 0.7rem;
        letter-spacing: 0.15em;
        border-bottom: 1px solid #1f2937;
        padding-bottom: 8px;
    }
    
    /* Text Content Styling */
    .content-text {
        color: #d1d5db;
        line-height: 1.8;
        font-size: 0.95rem;
    }

    /* Sidebar Button Refinement */
    .stButton>button {
        border-radius: 4px;
        border: 1px solid #374151;
        background-color: #1f2937;
        color: #e5e7eb;
        transition: all 0.3s ease;
    }
    .stButton>button:hover {
        border-color: #4b5563;
        background-color: #374151;
    }
    </style>
""", unsafe_allow_html=True)

st.markdown("<h1 class='main-title'>Multi Model RAG</h1>", unsafe_allow_html=True)

# 3. Session State Initialization
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())
    st.session_state.messages = []
    st.session_state.chat_title = "New Chat"
    st.session_state.db_ready = False

# 4. Directory Management
BASE_DIR = "temp_data"
USER_SESSION_DIR = os.path.join(BASE_DIR, st.session_state.session_id)
FILES_DIR = os.path.join(USER_SESSION_DIR, "files")
DB_DIR = os.path.join(USER_SESSION_DIR, "db")

os.makedirs(FILES_DIR, exist_ok=True)
os.makedirs(DB_DIR, exist_ok=True)

# 5. Helper Functions
def generate_document(messages):
    doc = Document()
    doc.add_heading('Chat Conversation Log', 0)
    for msg in messages:
        role = "User" if msg["role"] == "user" else f"AI ({msg.get('model_name', 'Unknown')})"
        p = doc.add_paragraph()
        runner = p.add_run(f"{role}:")
        runner.bold = True
        doc.add_paragraph(msg["content"])
        doc.add_paragraph("-" * 20)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

model_map = {
    "Llama 3.3 70B (Versatile)": "llama-3.3-70b-versatile",
    "Llama 3.1 8B (Instant)": "llama-3.1-8b-instant",
    "Llama 4 (Scout 17B)": "meta-llama/llama-4-scout-17b-16e-instruct",
    "Qwen 3 32B": "qwen/qwen3-32b",
    "GPT-OSS 20B": "openai/gpt-oss-20b"
}

@st.cache_resource
def get_rag_engine():
    return AdvancedRAG()

rag_engine = get_rag_engine()

# 6. Sidebar Implementation
with st.sidebar:
    st.header("Session Management")
    if st.button("New Chat", type="primary", use_container_width=True):
        if st.session_state.messages:
            st.session_state.chat_history.append({
                "id": st.session_state.session_id,
                "title": st.session_state.chat_title,
                "messages": st.session_state.messages,
                "db_ready": st.session_state.db_ready
            })
        st.session_state.session_id = str(uuid.uuid4())
        st.session_state.messages = []
        st.session_state.chat_title = "New Chat"
        st.session_state.db_ready = False
        st.rerun()

    st.markdown("---")
    st.header("Chat Archive")
    for chat in reversed(st.session_state.chat_history):
        if st.button(f"{chat['title']}", key=f"hist_{chat['id']}", use_container_width=True):
            if st.session_state.messages and st.session_state.session_id != chat['id']:
                 st.session_state.chat_history.append({
                    "id": st.session_state.session_id,
                    "title": st.session_state.chat_title,
                    "messages": st.session_state.messages,
                    "db_ready": st.session_state.db_ready
                })
            st.session_state.session_id = chat['id']
            st.session_state.messages = chat['messages']
            st.session_state.chat_title = chat['title']
            st.session_state.db_ready = chat['db_ready']
            st.session_state.chat_history = [c for c in st.session_state.chat_history if c['id'] != chat['id']]
            st.rerun()

    st.markdown("---")
    st.header("Data Export")
    if st.session_state.messages:
        docx_file = generate_document(st.session_state.messages)
        st.download_button(
            label="Download Session Log (DOCX)",
            data=docx_file,
            file_name=f"session_log_{st.session_state.session_id[:8]}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    st.markdown("---")
    st.header("System Settings")
    selected_model_friendly = st.selectbox("Intelligence Model", list(model_map.keys()), index=0)
    selected_model_id = model_map[selected_model_friendly]
    
    st.header("Document Processing")
    uploaded_files = st.file_uploader(
        "Upload Source Documents", 
        accept_multiple_files=True, 
        key=f"uploader_{st.session_state.session_id}"
    )
    
    if st.button("Initialize Knowledge Base", use_container_width=True):
        if uploaded_files:
            with st.spinner("Processing Corporate Data..."):
                if os.path.exists(FILES_DIR):
                    shutil.rmtree(FILES_DIR)
                os.makedirs(FILES_DIR)
                for file in uploaded_files:
                    file_path = os.path.join(FILES_DIR, file.name)
                    with open(file_path, "wb") as f:
                        f.write(file.getbuffer())
                status = rag_engine.process_documents(FILES_DIR, DB_DIR)
                if status == "Success":
                    st.success("Indexing Complete")
                    st.session_state.db_ready = True
                else:
                    st.error(f"Error: {status}")
        else:
            st.warning("Upload required documents.")

# 7. Chat Display Logic
for msg in st.session_state.messages:
    if msg["role"] == "user":
        st.markdown(f"""
            <div class="chat-container user-box">
                <div class="role-header">Inquiry</div>
                <div class="content-text">{msg['content']}</div>
            </div>
        """, unsafe_allow_html=True)
    else:
        model_info = msg.get('model_name', 'System Response')
        st.markdown(f"""
            <div class="chat-container ai-box">
                <div class="role-header">Response | {model_info}</div>
                <div class="content-text">{msg['content']}</div>
            </div>
        """, unsafe_allow_html=True)

# 8. Chat Input and Processing
if prompt := st.chat_input("Enter your inquiry..."):
    if not st.session_state.messages:
        st.session_state.chat_title = " ".join(prompt.split()[:5]) + "..."
    
    st.session_state.messages.append({"role": "user", "content": prompt})
    st.rerun()

if st.session_state.messages and st.session_state.messages[-1]["role"] == "user":
    if st.session_state.get("db_ready"):
        with st.spinner(f"Querying {selected_model_friendly}..."):
            response = rag_engine.query(
                query_text=st.session_state.messages[-1]["content"], 
                db_path=DB_DIR, 
                model_name=selected_model_id
            )
            st.session_state.messages.append({
                "role": "assistant", 
                "content": response,
                "model_name": selected_model_friendly
            })
            st.rerun()
    else:
        st.error("Knowledge base not initialized. Please upload documents.")